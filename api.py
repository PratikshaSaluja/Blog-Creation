from fastapi import FastAPI, Depends, HTTPException, status, UploadFile, File
from fastapi.middleware.cors import CORSMiddleware
from sqlalchemy.orm import Session
from sqlalchemy.exc import OperationalError
from sqlalchemy import text
from datetime import timedelta, date
import models, database, auth
from pydantic import BaseModel
from typing import Optional, List
import os
import re
from pathlib import Path
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse, StreamingResponse
import io
from fastapi.security import OAuth2PasswordBearer
import secrets
from datetime import datetime, timedelta

# Import the blog writer backend
from bwa_backend import app as blog_app


app = FastAPI(title="Blog Writing Agent API")

# Enable CORS for React frontend
origins = [
    "http://localhost:5173",
    "http://127.0.0.1:5173",
    "http://localhost:3000",
    "http://localhost:8000",
]

app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
    expose_headers=["Content-Disposition"],
)


# Create tables
models.Base.metadata.create_all(bind=database.engine)


class UserCreate(BaseModel):
    email: str
    password: str

class Token(BaseModel):
    access_token: str
    token_type: str

class ForgotPasswordRequest(BaseModel):
    email: str

class PasswordResetConfirm(BaseModel):
    email: str
    new_password: str

# Security
oauth2_scheme = OAuth2PasswordBearer(tokenUrl="login")

def get_current_user(db: Session = Depends(database.get_db), token: str = Depends(oauth2_scheme)):
    try:
        payload = auth.jwt.decode(token, auth.SECRET_KEY, algorithms=[auth.ALGORITHM])
        email: str = payload.get("sub")
        if email is None:
            raise HTTPException(status_code=401, detail="Invalid token")
    except Exception:
        raise HTTPException(status_code=401, detail="Invalid token")
        
    user = db.query(models.User).filter(models.User.email == email).first()
    if user is None:
        raise HTTPException(status_code=401, detail="User not found")
    return user

# --- Image Management ---

@app.get("/images/{filename}")
def get_image(filename: str, db: Session = Depends(database.get_db)):
    db_image = db.query(models.BlogImage).filter(models.BlogImage.filename == filename).first()
    if not db_image:
        raise HTTPException(status_code=404, detail="Image not found")
    
    ext = os.path.splitext(filename)[1].lower()
    media_type = "image/png"
    if ext == ".jpg" or ext == ".jpeg": media_type = "image/jpeg"
    elif ext == ".webp": media_type = "image/webp"
    elif ext == ".gif": media_type = "image/gif"
        
    return StreamingResponse(io.BytesIO(db_image.content), media_type=media_type)

@app.post("/blog/{filename}/upload-image")
async def upload_blog_image(filename: str, image: UploadFile = File(...), db: Session = Depends(database.get_db), current_user: models.User = Depends(get_current_user)):
    blog = db.query(models.Blog).filter(models.Blog.filename == filename, models.Blog.user_id == current_user.id).first()
    if not blog:
        raise HTTPException(status_code=404, detail="Blog not found")
    
    content = await image.read()
    image_filename = f"{secrets.token_hex(8)}_{image.filename}"
    
    db_image = models.BlogImage(
        filename=image_filename,
        content=content,
        blog_id=blog.id
    )
    db.add(db_image)
    db.commit()
    
    return {"url": f"/images/{image_filename}"}

@app.post("/signup", response_model=Token)
def signup(user: UserCreate, db: Session = Depends(database.get_db)):
    db_user = db.query(models.User).filter(models.User.email == user.email).first()
    if db_user:
        raise HTTPException(status_code=400, detail="Email already registered")
    
    hashed_password = auth.get_password_hash(user.password)
    new_user = models.User(email=user.email, hashed_password=hashed_password)
    db.add(new_user)
    db.commit()
    db.refresh(new_user)
    
    access_token = auth.create_access_token(data={"sub": new_user.email})
    return {"access_token": access_token, "token_type": "bearer"}

@app.post("/login", response_model=Token)
def login(user: UserCreate, db: Session = Depends(database.get_db)):
    db_user = db.query(models.User).filter(models.User.email == user.email).first()
    if not db_user or not auth.verify_password(user.password, db_user.hashed_password):
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Incorrect email or password",
            headers={"WWW-Authenticate": "Bearer"},
        )
    
    access_token = auth.create_access_token(data={"sub": db_user.email})
    return {"access_token": access_token, "token_type": "bearer"}

@app.get("/me")
def get_me(current_user: models.User = Depends(get_current_user)):
    return {"email": current_user.email, "id": current_user.id}

@app.post("/forgot-password")
def forgot_password(request: ForgotPasswordRequest, db: Session = Depends(database.get_db)):
    user = db.query(models.User).filter(models.User.email == request.email).first()
    if not user:
        raise HTTPException(status_code=404, detail="Email not found")
    return {"message": "Email validated"}

@app.post("/reset-password")
def reset_password(request: PasswordResetConfirm, db: Session = Depends(database.get_db)):
    # In this simplified flow, the token field in PasswordResetConfirm will be used for email if we don't want to change the model, 
    # but it's better to update the PasswordResetConfirm model to use 'email'.
    # Actually, I'll update the Pydantic model first.
    user = db.query(models.User).filter(models.User.email == request.email).first()
    
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    
    user.hashed_password = auth.get_password_hash(request.new_password)
    db.commit()
    
    return {"message": "Password successfully reset"}

# --- Blog Writer Endpoints ---

def safe_slug(title: str) -> str:
    s = title.strip().lower()
    s = re.sub(r"[^a-z0-9 _-]+", "", s)
    s = re.sub(r"\s+", "_", s).strip("_")
    return s or "blog"

@app.get("/past-blogs")
def get_past_blogs(db: Session = Depends(database.get_db), current_user: models.User = Depends(get_current_user)):
    db_blogs = db.query(models.Blog).filter(models.Blog.user_id == current_user.id).order_by(models.Blog.created_at.desc()).all()
    return [{"filename": b.filename, "title": b.title, "mtime": b.created_at.timestamp()} for b in db_blogs]

@app.get("/blog/{filename}")
def get_blog(filename: str, db: Session = Depends(database.get_db), current_user: models.User = Depends(get_current_user)):
    blog = db.query(models.Blog).filter(models.Blog.filename == filename, models.Blog.user_id == current_user.id).first()
    if not blog:
        raise HTTPException(status_code=404, detail="Blog not found")
    return {"content": blog.content}

class BlogUpdateRequest(BaseModel):
    content: str

@app.put("/blog/{filename}")
def update_blog(filename: str, request: BlogUpdateRequest, db: Session = Depends(database.get_db), current_user: models.User = Depends(get_current_user)):
    blog = db.query(models.Blog).filter(models.Blog.filename == filename, models.Blog.user_id == current_user.id).first()
    if not blog:
        raise HTTPException(status_code=404, detail="Blog not found")
    
    blog.content = request.content
    blog.updated_at = datetime.utcnow()
    db.commit()
    return {"message": "Blog updated successfully"}

class BlogGenerateRequest(BaseModel):
    topic: str
    as_of: Optional[str] = None

@app.post("/generate-blog")
async def generate_blog(request: BlogGenerateRequest, db: Session = Depends(database.get_db), current_user: models.User = Depends(get_current_user)):
    as_of = request.as_of or date.today().isoformat()
    
    inputs = {
        "topic": request.topic,
        "as_of": as_of,
        "recency_days": 7,
        "mode": "",
        "needs_research": False,
        "queries": [],
        "evidence": [],
        "plan": None,
        "sections": [],
        "merged_md": "",
        "md_with_placeholders": "",
        "image_specs": [],
        "generated_images": {},
        "final": "",
    }
    
    try:
        # Run the workflow
        result = blog_app.invoke(inputs)
        
        final_md = result.get("final", "")
        filename = None
        if final_md:
            from bwa_backend import _safe_slug
            title = result.get("topic")
            if not title and result.get("plan"):
                plan = result.get("plan")
                title = getattr(plan, "blog_title", None) if not isinstance(plan, dict) else plan.get("blog_title")
            
            filename = f"{_safe_slug(title or request.topic)}.md"
            
            # Save blog to database with retry for connection hiccups
            db_blog = None
            for attempt in range(3):
                try:
                    # Ping the database to ensure connection is alive
                    db.execute(text("SELECT 1"))
                    db_blog = db.query(models.Blog).filter(models.Blog.filename == filename, models.Blog.user_id == current_user.id).first()
                    break
                except (OperationalError, Exception) as e:
                    print(f"DB connection issue (attempt {attempt+1}): {e}")
                    db.rollback()
                    if attempt == 2:
                        raise
            
            if db_blog:
                db_blog.content = final_md
                db_blog.title = title or request.topic
            else:
                db_blog = models.Blog(
                    title=title or request.topic,
                    filename=filename,
                    content=final_md,
                    user_id=current_user.id
                )
                db.add(db_blog)
            
            if not db_blog:
                print("❌ Failed to save blog to database: db_blog is None after all retries")
                raise HTTPException(status_code=500, detail="Database error: Could not save blog")

            try:
                db.commit() # Commit to get db_blog.id
                db.refresh(db_blog)
            except Exception as e:
                db.rollback()
                # Final attempt to commit after rollback/refresh
                db.commit()
                db.refresh(db_blog)
            
            # Save images to database with similar caution
            generated_images = result.get("generated_images", {})
            for img_filename, img_data in generated_images.items():
                try:
                    db_image = db.query(models.BlogImage).filter(
                        models.BlogImage.filename == img_filename,
                        models.BlogImage.blog_id == db_blog.id
                    ).first()
                    if not db_image:
                        db_image = models.BlogImage(
                            filename=img_filename,
                            content=img_data,
                            blog_id=db_blog.id
                        )
                        db.add(db_image)
                    else:
                        db_image.content = img_data
                except Exception:
                    db.rollback()
                    # Re-try once
                    db.add(models.BlogImage(filename=img_filename, content=img_data, blog_id=db_blog.id))
            
            db.commit()
            
        # IMPORTANT: Remove binary data before returning to frontend. 
        # FastAPI's jsonable_encoder tries to .decode() bytes to UTF-8, which crashes for image data.
        result.pop("generated_images", None)
        result["filename"] = filename
        return result
    except Exception as e:
        db.rollback()
        print(f"Error generating blog: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/download-md/{filename}")
def download_md(filename: str, db: Session = Depends(database.get_db), current_user: models.User = Depends(get_current_user)):
    blog = db.query(models.Blog).filter(models.Blog.filename == filename, models.Blog.user_id == current_user.id).first()
    if not blog:
        raise HTTPException(status_code=404, detail="Blog not found")
    
    return StreamingResponse(
        io.BytesIO(blog.content.encode("utf-8")),
        media_type="text/markdown",
        headers={"Content-Disposition": f"attachment; filename={blog.filename}"}
    )

def markdown_to_docx(md_text: str, blog_title: str, db: Session, blog_id: int) -> bytes:
    """
    Converts markdown to a .docx file with embedded images fetched from DB.
    Uses a simplified regex-based parser.
    """
    from docx import Document
    from docx.shared import Inches
    import io

    doc = Document()
    doc.add_heading(blog_title, 0)

    # Simplified parser for headings, images, and text
    lines = md_text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue

        # Headings
        if line.startswith("# "):
            if line[2:].strip() != blog_title:
                doc.add_heading(line[2:].strip(), 1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), 2)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), 3)
        
        # Images: ![alt](src)
        elif "![" in line and "](" in line:
            matches = list(re.finditer(r"!\[(?P<alt>[^\]]*)\]\((?P<src>[^)]+)\)", line))
            if matches:
                last_end = 0
                for m in matches:
                    pre_text = line[last_end:m.start()].strip()
                    if pre_text:
                        doc.add_paragraph(pre_text)
                    
                    src = m.group("src").strip().lstrip("./")
                    # If it's a relative path starting with 'images/', fetch from DB
                    img_filename = src.split("/")[-1] if src.startswith("images/") else src
                    
                    db_image = db.query(models.BlogImage).filter(
                        models.BlogImage.filename == img_filename,
                        models.BlogImage.blog_id == blog_id
                    ).first()
                    
                    caption = None
                    if i + 1 < len(lines):
                        next_line = lines[i+1].strip()
                        if next_line.startswith("*") and next_line.endswith("*"):
                            caption = next_line[1:-1].strip()
                            i += 1 # skip caption line
                    
                    if db_image:
                        try:
                            from PIL import Image, features
                            img_buffer = io.BytesIO(db_image.content)
                            with Image.open(img_buffer) as img:
                                if img.mode in ("RGBA", "P"):
                                    img = img.convert("RGB")
                                img.save(img_buffer, format="JPEG", quality=85)
                            
                            img_buffer.seek(0)
                            doc.add_picture(img_buffer, width=Inches(5.5))
                            if caption:
                                doc.add_paragraph(caption, style='Caption')
                        except Exception as e:
                            error_type = type(e).__name__
                            print(f"❌ DOCX Image Error ({error_type}): {e}")
                            doc.add_paragraph(f"[Image Error ({error_type}): {e}]")
                    else:
                        doc.add_paragraph(f"[Image not found: {src}]")
                    
                    last_end = m.end()
                
                post_text = line[last_end:].strip()
                if post_text:
                    doc.add_paragraph(post_text)
            else:
                doc.add_paragraph(line)
        
        else:
            clean_line = re.sub(r"(?<!\!)\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", line)
            doc.add_paragraph(clean_line)
        
        i += 1

    out_io = io.BytesIO()
    doc.save(out_io)
    return out_io.getvalue()

@app.get("/download-docx/{filename}")
def download_docx(filename: str, db: Session = Depends(database.get_db), current_user: models.User = Depends(get_current_user)):
    blog = db.query(models.Blog).filter(models.Blog.filename == filename, models.Blog.user_id == current_user.id).first()
    if not blog:
        raise HTTPException(status_code=404, detail="Blog not found")
    
    docx_bytes = markdown_to_docx(blog.content, blog.title or blog.filename, db, blog.id)
    
    return StreamingResponse(
        io.BytesIO(docx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={safe_slug(blog.title or blog.filename)}.docx"}
    )


@app.get("/")
def health_check():
    return {"status": "ok", "message": "Blog Writing Agent API is running"}
